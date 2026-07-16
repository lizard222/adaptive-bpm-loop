# -*- coding: utf-8 -*-
"""Генерирует ПЛЕЙСХОЛДЕР-шаблон vkr_admission_order.docx (E1/T-36).

Реальных форм с кафедры нет (запрошены в T-02, не подтверждены) — этот
шаблон синтетический, нужен только чтобы проверить механизм подстановки
(docxtpl). Когда придут настоящие формы — файл заменяется, код агента не
меняется (agents/templater_agent.py работает с произвольным .docx с полями
{{ }}, не завязан на конкретные имена).

ВАЖНО (частая ошибка docxtpl): весь текст одного абзаца, включая {{ поле }},
должен идти ОДНИМ add_run(), иначе python-docx может разбить фигурные скобки
на разные XML-runs, и docxtpl не найдёт тег. Поэтому здесь — один run на
абзац, без пословного форматирования.

Перегенерировать: python templates/build_placeholder_template.py
"""
from pathlib import Path

from docx import Document

OUT = Path(__file__).parent / "vkr_admission_order.docx"

doc = Document()
doc.add_paragraph("ПРОЕКТ ПРИКАЗА (плейсхолдер-шаблон, не документ кафедры)")
doc.add_paragraph("")
doc.add_paragraph("Приказ № {{ order_number }} от {{ order_date }}")
doc.add_paragraph("")
doc.add_paragraph("О допуске к защите выпускной квалификационной работы")
doc.add_paragraph("")
doc.add_paragraph(
    "Допустить студента {{ student_name }} к защите выпускной "
    "квалификационной работы на тему «{{ topic }}»."
)
doc.add_paragraph("Научный руководитель: {{ supervisor_name }}.")
doc.add_paragraph("Дата защиты: {{ defense_date }}.")
doc.add_paragraph("")
doc.add_paragraph("Заведующий кафедрой ______________________")

doc.save(OUT)
print(f"Готово: {OUT}")
